from textblob import TextBlob
import requests
import json
import urllib
import random
import urllib2
import os
from docx import Document
from docx.shared import Inches




def download_web_image(url):
    name = random.randrange(1,1000)
    full_name = str(name) + ".jpg"
    urllib2.Request(url, full_name)
    urllib.urlretrieve(url, os.path.join(os.getcwd(), full_name))
    return full_name

# Different Online news articles
#s = "http://www.thehindu.com/todays-paper/tp-opinion/the-ai-battlefield/article20376166.ece"
#s = "https://timesofindia.indiatimes.com/india/no-country-can-thrive-without-equal-opportunity-for-half-its-population/articleshow/61827809.cms"
#s = "http://www.wionews.com/india-news/watch-india-an-inspiration-for-the-world-ivanka-trump-at-ges-2017-25369"
#s = "http://www.wionews.com/world/london-police-closes-roads-to-gherkin-skyscraper-after-suspicious-vehicle-found-25370"
#s = "https://timesofindia.indiatimes.com/city/hyderabad/made-in-india-bot-mitra-to-welcome-pm-narendra-modi-ivanka-trump-at-ges/articleshow/61827978.cms"
#s = "http://www.wionews.com/sports/cricket-australia-crush-england-by-10-wickets-in-1st-test-25214"
#s = "http://www.hindustantimes.com/business-news/india-gdp-can-grow-by-150bn-if-it-halves-gender-gap-ivanka-trump/story-ch3QRAcwZpCyGMSiZ3SPcO.html"
s = "http://www.hindustantimes.com/fashion-and-trends/manushi-chhillar-on-her-winning-moment-wish-i-had-given-a-more-lady-like-reaction/story-pCJlLA3yUeoz6QucVVZDPI.html?li_source=LI&li_medium=recommended-for-you"
s =  urllib.quote_plus(s)
f = "https://api.diffbot.com/v3/article?token="Your Diffbot API Key"&url=" + s


r = requests.get(f)
data = json.loads(r.content.decode("UTF-8"))
#print(data)
dd = data['objects'][0]['text']

#file_1 = open("C:\\Users\\chinkasoni\\Downloads\\pdftotext\\Freakonomics.txt","r")
#dd =  file_1.read()
#dd = dd.decode("UTF-8")
blob = TextBlob(dd)
#print(blob)
dh = blob.split(".")
print(len(dh))
#print(dh)
ll = []
fr = 0
qq = []
for lines in dh:
	blob = TextBlob(lines)
	qq = blob.tags
	for i in range(0,len(qq)):
		if (qq[i][1] == 'RB' or qq[i][1] == 'RBR' or qq[i][1] == 'RBS' or qq[i][1] == 'JJ' or qq[i][1] == 'JJS' or qq[i][1] == 'JJR' or qq[i][1] == 'NNP' or qq[i][1] == 'NNS'):
			fr+=1
	ll.append(fr)
	fr=0
print(data['objects'][0]['title'])
document = Document()

document.add_heading(data['objects'][0]['title'], 0)
#Checks to see if there is an image in the article. If there is, it will add it to the word document.
try:
        
        image = (data['objects'][0]['images'][0]['url'])
        ss = download_web_image(image)
        document.add_picture(ss, width=Inches(2.75), height=Inches(1.75))
except:
        p = document.add_paragraph("\n")
        p.add_run('Image Not Available.').italic = True


file_write = open("testfile4.txt","w")
file_write.write(data['objects'][0]['title'].encode("UTF-8"))
file_write.write("\n \n")
print("--------------")
kk=0
for i in range(0,len(ll)):
	if(ll[i]>8):
		file_write.write(dh[i].encode("UTF-8"))
		p = document.add_paragraph(dh[i])
		kk+=1
		
print(kk)

document.save('demo.docx')
file_write.close()
#This checks whether everthing was properly saved.
if(os.stat("testfile4.txt").st_size == 0):
        print("Nothing Happened")
else:
        print("Summarizing completed succesfully.")
