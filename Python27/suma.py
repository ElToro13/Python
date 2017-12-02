from summarize import datos
from docx import Document
from docx.shared import Inches
import os
import urllib
import random
import urllib2

def SaveInWord(zzz, filename="new"):
    zz = datos(zzz)

    TITLE, dd, LINK = zz.UserInput()
    lines, dh= zz.Summarize(dd)
    document = Document()

    document.add_heading(TITLE, 0)
    try:
        #print(LINK)    
        image = (LINK)
        ss = download_web_image(image)
        document.add_picture(ss, width=Inches(2.75), height=Inches(2.15))
    except Exception as e:
        print(e)
        p = document.add_paragraph("\n")
        p.add_run('Image Not Available.').italic = True


    #file_write = open("testfile4.txt","w")
    #file_write.write(data['objects'][0]['title'].encode("UTF-8"))
    #file_write.write("\n \n")
    print("--------------")
    kk=0

    for i in range(0,len(lines)):
        if(lines[i]>6):            #Cutoff
            p = document.add_paragraph(dh[i])
            
            kk+=1
                    
    print(kk)

    document.save("C:\\Users\\chinkasoni\\Desktop\\SMART WATCH\MLH\\"+filename+".docx")
    if(os.stat("C:\\Users\\chinkasoni\\Desktop\\SMART WATCH\\MLH\\"+filename+".docx").st_size == 0):
            print("Nothing Happened")
    else:
            print("Summarizing completed succesfully.")

def download_web_image(url):
        name = random.randrange(1,1000)
        full_name = str(name) + ".jpg"
        urllib2.Request(url, full_name)
        urllib.urlretrieve(url, os.path.join(os.getcwd(), full_name))
        return full_name


